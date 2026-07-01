from pathlib import Path
import re
import shutil
from typing import List, Dict, Any

import chromadb
from chromadb.utils import embedding_functions
from docx import Document
from PIL import Image
import pytesseract


# =========================
# 1. Project paths
# =========================

PROJECT_ROOT = Path(__file__).resolve().parents[1]

RAW_DOCUMENTS_DIR = PROJECT_ROOT / "data" / "raw" / "documents"
RAW_TEXT_DIR = PROJECT_ROOT / "data" / "raw" / "text"
RAW_IMAGES_DIR = PROJECT_ROOT / "data" / "raw" / "images"

EXTRACTED_TEXT_DIR = PROJECT_ROOT / "data" / "processed" / "extracted_text"

VECTOR_DB_DIR = PROJECT_ROOT / "vector_db" / "chroma"

COLLECTION_NAME = "studymate_knowledge_base"


# =========================
# 2. Index settings
# =========================

# 适合中英文混合资料的 embedding model
EMBEDDING_MODEL_NAME = "paraphrase-multilingual-MiniLM-L12-v2"

# 每个 chunk 的字符长度
CHUNK_SIZE = 1200

# chunk 之间保留一点重叠，避免上下文断裂
CHUNK_OVERLAP = 200

# 每次运行是否重新建立数据库
RESET_INDEX = True


# =========================
# 3. Utility functions
# =========================

def setup_tesseract_if_needed() -> None:
    """
    如果系统 PATH 里找不到 tesseract，就尝试使用 Windows 默认安装路径。
    """
    if shutil.which("tesseract"):
        return

    default_path = r"C:\Program Files\Tesseract-OCR\tesseract.exe"

    if Path(default_path).exists():
        pytesseract.pytesseract.tesseract_cmd = default_path
    else:
        raise FileNotFoundError(
            "Cannot find tesseract.exe. Please install Tesseract OCR "
            "or add it to system PATH."
        )


def read_text_file(file_path: Path) -> str:
    """
    读取 txt 文件。
    优先使用 utf-8，如果失败就用 gbk。
    """
    try:
        return file_path.read_text(encoding="utf-8")
    except UnicodeDecodeError:
        return file_path.read_text(encoding="gbk", errors="ignore")


def extract_docx_text(file_path: Path) -> str:
    """
    读取 Word .docx 文件中的文字。
    包括普通段落和表格内容。
    """
    doc = Document(file_path)
    parts = []

    # 读取普通段落
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            parts.append(text)

    # 读取表格内容
    for table in doc.tables:
        for row in table.rows:
            row_text = []
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:
                    row_text.append(cell_text)
            if row_text:
                parts.append(" | ".join(row_text))

    return "\n".join(parts)


def ocr_image(file_path: Path) -> str:
    """
    对图片进行 OCR。
    优先识别英文 + 简体中文。
    如果没有中文语言包，则自动退回英文。
    """
    setup_tesseract_if_needed()

    image = Image.open(file_path)

    try:
        text = pytesseract.image_to_string(image, lang="eng+chi_sim")
    except pytesseract.TesseractError:
        text = pytesseract.image_to_string(image, lang="eng")

    return text.strip()


def chunk_text(text: str, chunk_size: int = CHUNK_SIZE, overlap: int = CHUNK_OVERLAP) -> List[str]:
    """
    把长文本切成多个 chunk。
    ChromaDB 检索时更适合检索短 chunk，而不是整篇长文档。
    """
    text = text.strip()

    if not text:
        return []

    if len(text) <= chunk_size:
        return [text]

    chunks = []
    start = 0

    while start < len(text):
        end = start + chunk_size
        chunk = text[start:end].strip()

        if chunk:
            chunks.append(chunk)

        start = end - overlap

        if start < 0:
            start = 0

    return chunks


def parse_pdf_page_info(text: str) -> Dict[str, str]:
    """
    从 pdf_processor.py 生成的 txt 中提取 SOURCE_FILE 和 PAGE_NUMBER。
    """
    source_file = "unknown_pdf"
    page_number = "unknown"

    source_match = re.search(r"SOURCE_FILE:\s*(.+)", text)
    page_match = re.search(r"PAGE_NUMBER:\s*(.+)", text)

    if source_match:
        source_file = source_match.group(1).strip()

    if page_match:
        page_number = page_match.group(1).strip()

    return {
        "source_file": source_file,
        "page_number": page_number
    }


def create_records_from_text(
    text: str,
    source_file: str,
    source_path: str,
    source_type: str,
    modality: str,
    page_number: str = ""
) -> List[Dict[str, Any]]:
    """
    把一个文件的文本内容变成多个 records。
    每个 record 会成为 ChromaDB 中的一条记录。
    """
    chunks = chunk_text(text)
    records = []

    for i, chunk in enumerate(chunks):
        record = {
            "text": chunk,
            "metadata": {
                "source_file": source_file,
                "source_path": source_path,
                "source_type": source_type,
                "modality": modality,
                "page_number": str(page_number),
                "chunk_index": str(i)
            }
        }
        records.append(record)

    return records


# =========================
# 4. Load different file types
# =========================

def load_processed_pdf_texts() -> List[Dict[str, Any]]:
    """
    读取 data/processed/extracted_text/ 中由 pdf_processor.py 生成的 txt。
    这些内容包括 PDF_TEXT 和 OCR_TEXT。
    """
    records = []

    if not EXTRACTED_TEXT_DIR.exists():
        print(f"[Warning] Folder not found: {EXTRACTED_TEXT_DIR}")
        return records

    txt_files = list(EXTRACTED_TEXT_DIR.glob("*.txt"))

    for file_path in txt_files:
        text = read_text_file(file_path)
        page_info = parse_pdf_page_info(text)

        file_records = create_records_from_text(
            text=text,
            source_file=page_info["source_file"],
            source_path=str(file_path),
            source_type="processed_pdf_page",
            modality="pdf_text_plus_ocr",
            page_number=page_info["page_number"]
        )

        records.extend(file_records)

    print(f"Loaded processed PDF page text files: {len(txt_files)}")
    return records


def load_raw_txt_files() -> List[Dict[str, Any]]:
    """
    读取 data/raw/text/ 里的普通 txt 笔记。
    """
    records = []

    if not RAW_TEXT_DIR.exists():
        print(f"[Warning] Folder not found: {RAW_TEXT_DIR}")
        return records

    txt_files = list(RAW_TEXT_DIR.glob("*.txt"))

    for file_path in txt_files:
        text = read_text_file(file_path)

        file_records = create_records_from_text(
            text=text,
            source_file=file_path.name,
            source_path=str(file_path),
            source_type="raw_txt_note",
            modality="text"
        )

        records.extend(file_records)

    print(f"Loaded raw txt files: {len(txt_files)}")
    return records


def load_docx_files() -> List[Dict[str, Any]]:
    """
    读取 data/raw/documents/ 里的 Word .docx 文件。
    """
    records = []

    if not RAW_DOCUMENTS_DIR.exists():
        print(f"[Warning] Folder not found: {RAW_DOCUMENTS_DIR}")
        return records

    docx_files = list(RAW_DOCUMENTS_DIR.glob("*.docx"))

    for file_path in docx_files:
        text = extract_docx_text(file_path)

        file_records = create_records_from_text(
            text=text,
            source_file=file_path.name,
            source_path=str(file_path),
            source_type="word_document",
            modality="text"
        )

        records.extend(file_records)

    print(f"Loaded docx files: {len(docx_files)}")
    return records


def load_image_files() -> List[Dict[str, Any]]:
    """
    读取 data/raw/images/ 里的图片，并进行 OCR。
    这部分是额外图片，不是 PDF 自动生成的页面图。
    """
    records = []

    if not RAW_IMAGES_DIR.exists():
        print(f"[Warning] Folder not found: {RAW_IMAGES_DIR}")
        return records

    image_extensions = [".png", ".jpg", ".jpeg", ".bmp", ".tif", ".tiff"]

    image_files = [
        file_path
        for file_path in RAW_IMAGES_DIR.iterdir()
        if file_path.suffix.lower() in image_extensions
    ]

    for file_path in image_files:
        try:
            text = ocr_image(file_path)
        except Exception as e:
            print(f"[Warning] OCR failed for {file_path.name}: {e}")
            continue

        if not text.strip():
            print(f"[Warning] No OCR text found in image: {file_path.name}")
            continue

        combined_text = f"""
SOURCE_FILE: {file_path.name}
IMAGE_PATH: {file_path}

OCR_TEXT:
{text}
""".strip()

        file_records = create_records_from_text(
            text=combined_text,
            source_file=file_path.name,
            source_path=str(file_path),
            source_type="raw_image_ocr",
            modality="image_ocr"
        )

        records.extend(file_records)

    print(f"Loaded image files with OCR: {len(image_files)}")
    return records


def load_all_records() -> List[Dict[str, Any]]:
    """
    读取所有知识库内容：
    1. PDF 处理后的 page txt
    2. 普通 txt
    3. Word docx
    4. 图片 OCR
    """
    all_records = []

    all_records.extend(load_processed_pdf_texts())
    all_records.extend(load_raw_txt_files())
    all_records.extend(load_docx_files())
    all_records.extend(load_image_files())

    return all_records


# =========================
# 5. Build ChromaDB index
# =========================

def build_chroma_index(records: List[Dict[str, Any]]) -> None:
    """
    把所有 records 写入 ChromaDB。
    """
    if not records:
        print("No records found. Please add documents or run pdf_processor.py first.")
        return

    if RESET_INDEX and VECTOR_DB_DIR.exists():
        shutil.rmtree(VECTOR_DB_DIR)

    VECTOR_DB_DIR.mkdir(parents=True, exist_ok=True)

    embedding_function = embedding_functions.SentenceTransformerEmbeddingFunction(
        model_name=EMBEDDING_MODEL_NAME
    )

    client = chromadb.PersistentClient(path=str(VECTOR_DB_DIR))

    collection = client.get_or_create_collection(
        name=COLLECTION_NAME,
        embedding_function=embedding_function
    )

    documents = []
    metadatas = []
    ids = []

    for idx, record in enumerate(records):
        documents.append(record["text"])
        metadatas.append(record["metadata"])
        ids.append(f"doc_{idx}")

    batch_size = 100

    for start in range(0, len(documents), batch_size):
        end = start + batch_size

        collection.add(
            documents=documents[start:end],
            metadatas=metadatas[start:end],
            ids=ids[start:end]
        )

        print(f"Indexed chunks {start + 1} to {min(end, len(documents))}")

    print("\nChromaDB index built successfully.")
    print(f"Total chunks indexed: {len(documents)}")
    print(f"Vector DB path: {VECTOR_DB_DIR}")


def test_search() -> None:
    """
    简单测试向量数据库是否可以检索。
    """
    embedding_function = embedding_functions.SentenceTransformerEmbeddingFunction(
        model_name=EMBEDDING_MODEL_NAME
    )

    client = chromadb.PersistentClient(path=str(VECTOR_DB_DIR))

    collection = client.get_collection(
        name=COLLECTION_NAME,
        embedding_function=embedding_function
    )

    query = "What is retrieval augmented generation?"

    results = collection.query(
        query_texts=[query],
        n_results=3
    )

    print("\nTest search query:")
    print(query)

    print("\nTop results:")

    for i in range(len(results["documents"][0])):
        metadata = results["metadatas"][0][i]
        document = results["documents"][0][i]

        print("\n----------------------------")
        print(f"Rank: {i + 1}")
        print(f"Source file: {metadata.get('source_file')}")
        print(f"Page number: {metadata.get('page_number')}")
        print(f"Modality: {metadata.get('modality')}")
        print("Text preview:")
        print(document[:500])


def main() -> None:
    print("Building StudyMate ChromaDB index...\n")

    records = load_all_records()

    print(f"\nTotal records before indexing: {len(records)}")

    build_chroma_index(records)

    if records:
        test_search()


if __name__ == "__main__":
    main()